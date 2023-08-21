import docx
from collections import Counter
import numpy as np
import json

# only three outliers, just hard code since the list wont change.
def clean_outliers(word):
    outlier_replacements = {
        "Māori, māori": "māori",
        "i nāianei (written as ināianei in older texts)": "i nāianei",
        "haramai (haere mai)": "haere mai"
    }
    return outlier_replacements.get(word, word)

# Read docx file containing maori word frequency, create list of words from table.
def extract_table_data(docx_path, table_index=0):
    doc = docx.Document(docx_path)
    
    if table_index >= len(doc.tables):
        print("Table index out of range")
        return []
    
    table = doc.tables[table_index]
    table_data = []
    
    for row in table.rows:
        word = row.cells[0].text.strip()
        cleaned_word = clean_outliers(word)
        table_data.append(cleaned_word)
    
    return table_data


def apply_zipfs_law(word_list):
    # Step 1: Calculate frequencies of each word
    word_frequencies = Counter(word_list)
    
    # Step 2: Apply Zipf's Law to calculate relative frequencies
    sorted_word_frequencies = sorted(word_frequencies.values(), reverse=True)
    relative_frequencies = [1 / (i + 1) for i in range(len(sorted_word_frequencies))]
    
    # Step 3: Normalize relative frequencies
    total_relative_frequency = sum(relative_frequencies)
    normalized_probabilities = [freq / total_relative_frequency for freq in relative_frequencies]
    
    # Create a dictionary with words and their probabilities
    word_probabilities = {word: prob for word, prob in zip(word_frequencies.keys(), normalized_probabilities)}
    
    return word_probabilities

if __name__ == "__main__":
    docx_path = "../raw_data/maori_word_frequency.docx"  # Replace with the actual path to your Word document
    table_index = 0  # Index of the table you want to extract
    
    table_data = extract_table_data(docx_path, table_index)
    
    output_file_path = "../clean_data/maori_word_frequency.json"
    
    probabilities_dict = apply_zipfs_law(table_data)
    
    with open(output_file_path, "w", encoding='utf-8') as output_file:
        json.dump(probabilities_dict, output_file, indent=4)

    print("Data written to", output_file_path)
