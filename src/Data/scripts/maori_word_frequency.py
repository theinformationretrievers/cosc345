import docx
from collections import Counter
import json
import argparse


def main():
    DOCX_PATH = "../raw_data/maori_word_frequency.docx"
    TABLE_INDX = 0  # Index of the table to extract

    output_file_path = handle_args()

    table_data = extract_table_data(DOCX_PATH, TABLE_INDX)
    probabilities_dict = apply_zipfs_law(table_data)

    with open(output_file_path, "w", encoding='utf-8') as output_file:
        json.dump(probabilities_dict, output_file, indent=4)

    print("Data written to", output_file_path)


def clean_outliers(word):
    """
    Cleans outliers in Maori words using predefined replacements. Hard coded as list won't change

    Args:
        word (str): The Maori word to be cleaned.

    Returns:
        str: The cleaned Maori word, with outliers replaced if applicable.

    Example:
        >>> clean_outliers("i nāianei (written as ināianei in older texts)")
        'i nāianei'
        >>> clean_outliers("whānau")
        'whānau'  # No outlier match, original word returned
    """
    outlier_replacements = {
        "Māori, māori": "māori",
        "i nāianei (written as ināianei in older texts)": "i nāianei",
        "haramai (haere mai)": "haere mai"
    }
    return outlier_replacements.get(word, word)


def extract_table_data(docx_path, table_index=0):
    """
    Extracts data from a specific table in a DOCX file containing Maori word frequencies.

    This function reads a DOCX file and extracts data from a specified table, assumed to contain
    Maori word frequencies. The extracted words are cleaned using the `clean_outliers` function
    to ensure consistent formatting.

    Args:
        docx_path (str): The path to the DOCX file.
        table_index (int, optional): The index of the table to extract data from. Default is 0.

    Returns:
        list: A list of cleaned Maori words extracted from the specified table.

    Example:
        >>> docx_path = 'maori_frequencies.docx'
        >>> words = extract_table_data(docx_path, table_index=1)
        >>> print(words)
        ['whānau', 'tamariki', 'marae', 'kai', 'iwi']

    """
    doc = docx.Document(docx_path)

    if table_index >= len(doc.tables):
        print("Table index out of range")
        return []

    table = doc.tables[table_index]
    table_data = []

    for row in table.rows:
        word = row.cells[0].text.strip()
        cleaned_word = clean_outliers(word)
        table_data.append(cleaned_word)

    return table_data


def apply_zipfs_law(word_list):
    """
    Applies Zipf's Law to a list of words and calculates their probabilities.

    Args:
        word_list (list): A list of words for which to calculate probabilities.

    Returns:
        dict: A dictionary containing words as keys and their corresponding
              calculated probabilities as values.

    Example:
        >>> words = ['apple', 'banana', 'apple', 'orange', 'banana', 'apple']
        >>> probabilities = apply_zipfs_law(words)
        >>> print(probabilities)
        {'apple': 0.5, 'banana': 0.3333333333333333, 'orange': 0.16666666666666666}
    """
    # Step 1: Calculate frequencies of each word
    word_frequencies = Counter(word_list)

    # Step 2: Apply Zipf's Law to calculate relative frequencies
    sorted_word_frequencies = sorted(word_frequencies.values(), reverse=True)
    relative_frequencies = [1 / (i + 1)
                            for i in range(len(sorted_word_frequencies))]

    # Step 3: Normalize relative frequencies
    total_relative_frequency = sum(relative_frequencies)
    normalized_probabilities = [
        freq / total_relative_frequency for freq in relative_frequencies]

    # Create a dictionary with words and their probabilities
    word_probabilities = {word: prob for word, prob in zip(
        word_frequencies.keys(), normalized_probabilities)}

    return word_probabilities


def handle_args():
    """
    Handles command-line arguments for custom output file paths.

    This function uses the argparse library to handle the --filepath argument, allowing users
    to specify a custom output file path for saving the generated JSON data. If no file path
    is provided, a default output path is used.

    Returns:
        str: The selected or default output file path.
    """
    parser = argparse.ArgumentParser(
        description="Input filepath to describe custom output, otherwise placed in clean_data.")
    parser.add_argument('--filepath', help='Path to the file')
    args = parser.parse_args()

    if args.filepath is None:
        output_file_path = "../clean_data/maori_frequencies.json"
    else:
        output_file_path = args.filepath

    return output_file_path


if __name__ == "__main__":
    main()
